import docx
import pandas as pd

def writeContracts():
    """
    Returns a docx for every person in the csv.
    """
    data = pd.read_csv("ContractInfoSP21.csv")
    for ind in data.index:

        getDate= data['Date'][ind]
        getAddress= str(data['Address'][ind])
        getAddress = getAddress.replace('-', '\n')
        getName= data['Name'][ind]
        getStart= data['Start'][ind]
        getEnd= data['End'][ind]
        getTeam= data['Team'][ind]
        getExpir = data['Expiration'][ind]
        getDir = data['Director'][ind]

        workData = {
            "{ADDRESS}": getAddress,
            "{NAME}": getName,
            "{START}": getStart,
            "{END}": getEnd,
            "{TEAM}": getTeam,
            "{EXPIRATION}": getExpir,
            "{TODAY}": getDate
        }
        if getDir == 'Y':
            doc = docx.Document("./Templates/Template" + getTeam + "Dir.docx")

        else:
            doc = docx.Document("./Templates/Template" + getTeam +  ".docx")


        for item in workData.keys():
            for para in doc.paragraphs:
                para.text = para.text.replace(item, workData[item])

        doc.save("./Ready/" + getName + " PoliciContract.docx")

if __name__ == "__main__":
    writeContracts()
